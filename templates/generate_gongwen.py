#!/usr/bin/env python3
"""
政企公文 Word 文档生成器
基于 GB/T 9704-2012《党政机关公文格式》标准
依赖: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime


class GongwenGenerator:
    """GB/T 9704-2012 标准公文生成器"""

    # 标准颜色
    RED = RGBColor(0xFF, 0x00, 0x00)
    BLACK = RGBColor(0x00, 0x00, 0x00)

    # 标准字号映射 (name -> pt)
    FONT_SIZES = {
        "初号": 42, "小初": 36,
        "一号": 26, "小一": 24,
        "二号": 22, "小二": 18,
        "三号": 16, "小三": 15,
        "四号": 14, "小四": 12,
        "五号": 10.5, "小五": 9,
    }

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = Document()
        self._setup_page()
        self._setup_styles()
        self._header_set = False
        self._body_started = False
        self._cc_added = False
        self._print_info_added = False

    # ── 页面设置 ──────────────────────────────────

    def _setup_page(self):
        """设置页面：A4纸，GB/T 9704标准页边距"""
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)      # 上 37mm
        section.bottom_margin = Cm(3.5)   # 下 35mm
        section.left_margin = Cm(2.8)     # 左 28mm
        section.right_margin = Cm(2.6)    # 右 26mm
        # 页眉页脚距边界
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)

    def _setup_styles(self):
        """设置默认正文样式：仿宋3号，固定行距28.95pt"""
        style = self.doc.styles['Normal']
        self._set_style_font(style, 'FangSong', '仿宋', Pt(16))
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(28.95)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    # ── 字体工具 ──────────────────────────────────

    @staticmethod
    def _set_run_font(run, ascii_font='FangSong', east_asia_font='仿宋',
                      size=None, bold=False, color=None):
        """设置 run 的中西文字体"""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), ascii_font)
        rFonts.set(qn('w:hAnsi'), ascii_font)
        rFonts.set(qn('w:eastAsia'), east_asia_font)
        if size:
            run.font.size = size
        if bold:
            run.font.bold = True
        if color:
            run.font.color.rgb = color

    @staticmethod
    def _set_style_font(style, ascii_font, east_asia_font, size):
        """设置样式的中西文字体"""
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), ascii_font)
        rFonts.set(qn('w:hAnsi'), ascii_font)
        rFonts.set(qn('w:eastAsia'), east_asia_font)
        style.font.size = size

    def _add_paragraph(self, text, ascii_font='FangSong', east_asia_font='仿宋',
                       size=Pt(16), bold=False, color=None,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       first_line_indent=None, right_indent=None,
                       space_before=None, space_after=None,
                       line_spacing=None):
        """添加一个格式化段落"""
        para = self.doc.add_paragraph()
        para.alignment = alignment
        pf = para.paragraph_format
        if first_line_indent is not None:
            pf.first_line_indent = first_line_indent
        if right_indent is not None:
            pf.right_indent = right_indent
        if space_before is not None:
            pf.space_before = space_before
        if space_after is not None:
            pf.space_after = space_after
        if line_spacing is not None:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = line_spacing
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(28.95)

        run = para.add_run(text)
        self._set_run_font(run, ascii_font, east_asia_font, size, bold, color)
        return para

    # ── 版头（红头）──────────────────────────────

    def set_header(self, org_name: str, doc_number: str = "",
                   secrecy: str = "", urgency: str = ""):
        """
        设置版头（红头区域）
        Args:
            org_name: 发文机关名称，如 "国务院办公厅"
            doc_number: 发文字号，如 "国办发〔2026〕15号"
            secrecy: 密级，如 "秘密"、"机密"、"绝密"（可选）
            urgency: 紧急程度，如 "特急"、"加急"（可选）
        """
        # 份号（如有密级则显示）
        if secrecy:
            self._add_paragraph(
                "                          000001",
                size=Pt(16), space_before=Pt(0), space_after=Pt(0)
            )

        # 密级和紧急程度
        if secrecy or urgency:
            label = f"{'　' * 2}{secrecy}{'　' * 4}{urgency}" if secrecy and urgency \
                else f"{'　' * 2}{secrecy}{urgency}"
            self._add_paragraph(
                label, ascii_font='SimHei', east_asia_font='黑体',
                size=Pt(16), color=self.RED,
                space_before=Pt(0), space_after=Pt(6)
            )

        # 发文机关标志（红色大字，居中）
        self._add_paragraph(
            org_name,
            ascii_font='STZhongsong', east_asia_font='华文中宋',
            size=Pt(26), bold=True, color=self.RED,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(12), space_after=Pt(0)
        )

        # 版头红色分隔线
        self._add_red_line()

        # 发文字号（居中）
        if doc_number:
            self._add_paragraph(
                doc_number,
                size=Pt(16),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(6), space_after=Pt(6)
            )

        # 正文开始红线
        self._add_red_line()
        self._header_set = True

    def _add_red_line(self):
        """添加红色分隔线（用段落下边框模拟）"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = para.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(6)

        # 段落下边框设为红色
        pPr = para._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="FF0000"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    # ── 标题 ─────────────────────────────────────

    def set_title(self, title: str):
        """设置公文标题（小标宋2号，居中）"""
        self._add_paragraph(
            title,
            ascii_font='STZhongsong', east_asia_font='华文中宋',
            size=Pt(22), bold=True, color=self.RED,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(16), space_after=Pt(8),
            line_spacing=Pt(32)
        )

    def set_main_recipient(self, recipient: str):
        """设置主送机关（仿宋3号，左顶格）"""
        self._add_paragraph(
            recipient,
            size=Pt(16),
            space_before=Pt(8), space_after=Pt(0),
            first_line_indent=Pt(0)
        )

    # ── 正文 ─────────────────────────────────────

    def add_body(self, text: str, level: int = 0):
        """
        添加正文内容
        Args:
            text: 文本内容
            level: 标题层级
                   0 = 普通正文段落（仿宋3号，首行缩进2字符）
                   1 = 一级标题（黑体3号，如 "一、总体要求"）
                   2 = 二级标题（楷体3号，如 "（一）基本原则"）
                   3 = 三级标题（仿宋加粗3号，如 "1. 主要目标"）
                   4 = 四级标题（仿宋3号，如 "（1）具体内容"）
        """
        self._body_started = True

        if level == 1:
            # 一级标题：黑体
            self._add_paragraph(
                text, ascii_font='SimHei', east_asia_font='黑体',
                size=Pt(16), bold=True,
                space_before=Pt(8), space_after=Pt(0),
                first_line_indent=Pt(32)  # 2字符
            )
        elif level == 2:
            # 二级标题：楷体
            self._add_paragraph(
                text, ascii_font='KaiTi', east_asia_font='楷体',
                size=Pt(16), bold=False,
                space_before=Pt(4), space_after=Pt(0),
                first_line_indent=Pt(32)
            )
        elif level == 3:
            # 三级标题：仿宋加粗
            self._add_paragraph(
                text, size=Pt(16), bold=True,
                space_before=Pt(4), space_after=Pt(0),
                first_line_indent=Pt(32)
            )
        elif level == 4:
            # 四级标题：仿宋
            self._add_paragraph(
                text, size=Pt(16),
                space_before=Pt(2), space_after=Pt(0),
                first_line_indent=Pt(32)
            )
        else:
            # 普通正文段落：仿宋3号，首行缩进2字符
            self._add_paragraph(
                text, size=Pt(16),
                first_line_indent=Pt(32)  # 16pt × 2 = 32pt
            )

    # ── 附件 ─────────────────────────────────────

    def add_attachment_note(self, note: str):
        """添加附件说明（仿宋3号，左空2字）"""
        self._add_paragraph(
            "",
            space_before=Pt(8)  # 附件说明前空一行
        )
        self._add_paragraph(
            "附件：" + note if not note.startswith("附件") else note,
            size=Pt(16),
            first_line_indent=Pt(32)
        )

    # ── 落款 ─────────────────────────────────────

    def set_signature(self, org_name: str, date: str = ""):
        """
        设置落款：发文机关名称 + 成文日期
        右对齐，右空4字（约 Cm(2.8)）
        """
        # 空行间隔
        self._add_paragraph("", space_before=Pt(16))

        # 发文机关名称（右对齐，右空约4字符）
        self._add_paragraph(
            org_name,
            size=Pt(16),
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            right_indent=Cm(2.8)
        )

        # 成文日期
        if not date:
            date = datetime.datetime.now().strftime("%Y年%m月%d日")
        self._add_paragraph(
            date,
            size=Pt(16),
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            right_indent=Cm(2.8)
        )

    # ── 附注 ─────────────────────────────────────

    def add_annotation(self, text: str):
        """添加附注（仿宋3号，左空2字，加圆括号）"""
        if not text.startswith("（"):
            text = f"（{text}）"
        self._add_paragraph(
            text, size=Pt(16),
            first_line_indent=Pt(32)
        )

    # ── 抄送与印发 ───────────────────────────────

    def add_cc(self, cc_text: str):
        """添加抄送机关（仿宋4号）"""
        self._cc_added = True
        # 分隔线
        self._add_cc_line()
        self._add_paragraph(
            f"抄送：{cc_text}",
            size=Pt(14),  # 4号 = 14pt
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(4)
        )

    def _add_cc_line(self):
        """添加抄送分隔线"""
        para = self.doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(2)
        pf.line_spacing = Pt(6)
        pPr = para._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def add_print_info(self, print_org: str, print_date: str = ""):
        """添加印发机关和印发日期（仿宋4号）"""
        if not print_date:
            print_date = datetime.datetime.now().strftime("%Y年%m月%d日印发")

        self._add_paragraph(
            f"{print_org}{'　' * 4}{print_date}",
            size=Pt(14),
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(4)
        )
        self._print_info_added = True

    # ── 页码 ─────────────────────────────────────

    def _setup_page_number(self):
        """设置页码格式：—×—（居中，4号仿宋）"""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加"—"前缀
        run1 = para.add_run("— ")
        self._set_run_font(run1, 'FangSong', '仿宋', Pt(14))

        # 添加页码域代码
        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
            f'  <w:r>'
            f'    <w:rPr>'
            f'      <w:rFonts w:ascii="FangSong" w:hAnsi="FangSong" w:eastAsia="仿宋"/>'
            f'      <w:sz w:val="28"/>'  # 4号=14pt, val=半pt=28
            f'    </w:rPr>'
            f'    <w:t>1</w:t>'
            f'  </w:r>'
            f'</w:fldSimple>'
        )
        para._element.append(parse_xml(fld_xml))

        # 添加"—"后缀
        run2 = para.add_run(" —")
        self._set_run_font(run2, 'FangSong', '仿宋', Pt(14))

    # ── 保存 ─────────────────────────────────────

    def save(self):
        """保存文档"""
        self._setup_page_number()
        self.doc.save(self.filepath)
        print(f"✅ 公文已生成: {self.filepath}")


# ── 使用示例 ─────────────────────────────────────

def demo():
    """演示：生成一份标准通知公文"""
    gen = GongwenGenerator("示例公文_通知.docx")

    # 版头
    gen.set_header(
        org_name="国务院办公厅",
        doc_number="国办发〔2026〕15号",
        secrecy="",
        urgency=""
    )

    # 标题
    gen.set_title("关于加快推进数字化转型的通知")

    # 主送机关
    gen.set_main_recipient("各省、自治区、直辖市人民政府，国务院各部委、各直属机构：")

    # 正文
    gen.add_body("为深入贯彻落实党中央、国务院关于数字中国建设的决策部署，"
                 "加快推进各领域数字化转型，经国务院同意，现就有关事项通知如下：")

    gen.add_body("一、总体要求", level=1)

    gen.add_body("坚持以习近平新时代中国特色社会主义思想为指导，全面贯彻党的二十大精神，"
                 "完整、准确、全面贯彻新发展理念，加快构建新发展格局，"
                 "以推动高质量发展为主题，以深化供给侧结构性改革为主线，"
                 "以满足人民日益增长的美好生活需要为根本目的。")

    gen.add_body("（一）基本原则", level=2)

    gen.add_body("坚持统筹规划、分步实施。加强顶层设计和系统谋划，"
                 "明确数字化转型的目标、路径和重点任务，分阶段、分步骤有序推进。")

    gen.add_body("（二）发展目标", level=2)

    gen.add_body("到2027年，数字化转型取得显著成效。数字经济核心产业增加值占GDP比重"
                 "达到12%以上，关键业务环节全面数字化的企业比例达到65%以上。")

    gen.add_body("1. 近期目标", level=3)
    gen.add_body("到2025年底，各重点领域数字化转型实施方案全部出台，"
                 "示范引领作用初步显现。")

    gen.add_body("2. 远期目标", level=3)
    gen.add_body("到2030年，数字化转型整体迈上新台阶，"
                 "为全面建设社会主义现代化国家提供有力支撑。")

    gen.add_body("二、重点任务", level=1)

    gen.add_body("（一）推进产业数字化转型", level=2)
    gen.add_body("加快制造业数字化转型步伐。深入实施智能制造工程，"
                 "推动工业互联网创新发展，培育一批数字化转型标杆企业。")

    gen.add_body("（二）加快数字产业化发展", level=2)
    gen.add_body("做强做优做大数字经济。培育壮大人工智能、大数据、"
                 "云计算、区块链等新兴数字产业。")

    gen.add_body("三、保障措施", level=1)

    gen.add_body("各地区、各部门要高度重视数字化转型工作，"
                 "加强组织领导，明确责任分工，确保各项任务落到实处。"
                 "国务院有关部门要加强统筹协调和督促检查，"
                 "重大情况及时报告国务院。")

    # 附件
    gen.add_attachment_note("1. 数字化转型重点任务分工方案")

    # 落款
    gen.set_signature("国务院办公厅", "2026年5月26日")

    # 附注
    gen.add_annotation("此件公开发布")

    # 抄送
    gen.add_cc("中央办公厅，全国人大常委会办公厅，全国政协办公厅，"
               "国务院各部门，各民主党派中央。")

    # 印发
    gen.add_print_info("国务院办公厅秘书局", "2026年5月27日印发")

    gen.save()


if __name__ == "__main__":
    demo()
